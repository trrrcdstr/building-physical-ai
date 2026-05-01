"""读取DOCX中的VR链接"""
import docx, os, re, json

def anonymize(text):
    if not text:
        return ""
    t = str(text)
    t = re.sub(r"龙湖|Longfor", "[某房企]", t)
    t = re.sub(r"[一-龥]{2,6}·[一-龥]{2,10}", "[项目名]", t)
    return t.strip()

docx_path = r"C:\Users\Administrator\Desktop\27套3D效果图链接.docx"
try:
    doc = docx.Document(docx_path)
    print("段落数:", len(doc.paragraphs))
    print("表格数:", len(doc.tables))
    
    all_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            all_text.append(anonymize(p.text))
    
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(anonymize(cell.text))
    
    print("\n=== 全部文本 ===")
    for line in all_text:
        print(line)
    
    # 保存清洗后的文本
    out_dir = r"C:\Users\Administrator\.qclaw\workspace\projects\building-physical-ai\data\processed"
    with open(os.path.join(out_dir, "vr_links_raw.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(all_text))
    
    # 提取VR链接（http/https）
    vr_links = []
    for line in all_text:
        urls = re.findall(r"https?://[^\s，、。]+", line)
        vr_links.extend(urls)
    
    print("\n=== VR链接 ===")
    for u in vr_links:
        print(u)
    
    print("\n共找到 {} 个链接".format(len(vr_links)))
    
except Exception as e:
    print("Error:", e)
    import traceback; traceback.print_exc()
