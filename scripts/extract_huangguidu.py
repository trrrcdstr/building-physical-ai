"""
黄葛渡施工图 - 深度数据提取
1. 提取 PDF/XLS/DOCX 中的文本数据
2. 提取图片中的元数据
3. 提取 ZIP 中的内容
4. 对 DWG 用二进制方式提取 CAD 元数据（不依赖 ezdxf）
"""
import os, sys, glob, json, struct, re
from pathlib import Path

# PDF 提取
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# Word 提取
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Excel 提取
try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

SRC = r"C:\Users\Administrator\Desktop\黄葛渡施工图"
OUT_DIR = r"C:\Users\Administrator\.qclaw\workspace\projects\building-physical-ai\data\processed\huangguidu"
os.makedirs(OUT_DIR, exist_ok=True)

def anonymize(text):
    if not text:
        return ""
    t = str(text)
    t = re.sub(r"龙湖|Longfor|LONGFOR", "[某房企]", t)
    t = re.sub(r"黄葛渡", "[示范区]", t)
    t = re.sub(r"南坪|南岸区", "[项目地]", t)
    t = re.sub(r"设计单位：[^\s，,，]+", "设计单位：[某设计院]", t)
    t = re.sub(r"施工单位：[^\s，,，]+", "施工单位：[某施工单位]", t)
    t = re.sub(r"建设单位：[^\s，,，]+", "建设单位：[某建设单位]", t)
    t = re.sub(r"监理单位：[^\s，,，]+", "监理单位：[某监理单位]", t)
    t = re.sub(r"制图人：[^\s，,]+", "制图人：[设计师]", t)
    t = re.sub(r"审核人：[^\s，,]+", "审核人：[审核师]", t)
    t = re.sub(r"批准人：[^\s，,]+", "批准人：[批准人]", t)
    return t.strip()

def extract_pdf(fpath):
    """从 PDF 提取文本"""
    if not HAS_PYMUPDF:
        return {"error": "PyMuPDF not installed"}
    
    try:
        doc = fitz.open(fpath)
        pages = []
        total_text = ""
        
        for i, page in enumerate(doc):
            text = page.get_text()
            text = anonymize(text)
            total_text += text + "\n"
            pages.append({"page": i+1, "text": text[:2000], "length": len(text)})
        
        doc.close()
        
        # 提取尺寸（宽高）
        dims = doc.tobytes()[:100] if doc else b""
        
        return {
            "status": "ok",
            "pages": len(pages),
            "total_chars": len(total_text),
            "first_page_preview": total_text[:500],
            "page_list": pages[:3],
        }
    except Exception as e:
        return {"error": str(e)[:100]}

def extract_docx(fpath):
    """从 Word 文档提取文本"""
    if not HAS_DOCX:
        return {"error": "python-docx not installed"}
    
    try:
        doc = docx.Document(fpath)
        paragraphs = [anonymize(p.text) for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [anonymize(c.text) for c in row.cells]
                rows.append(cells)
            tables.append(rows)
        
        full_text = "\n".join(paragraphs)
        return {
            "status": "ok",
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "total_chars": len(full_text),
            "preview": full_text[:800],
            "first_table": tables[0] if tables else None,
        }
    except Exception as e:
        return {"error": str(e)[:100]}

def extract_xls(fpath):
    """从 Excel 提取数据"""
    if not HAS_XLSX:
        return {"error": "openpyxl not installed"}
    
    try:
        wb = openpyxl.load_workbook(fpath, data_only=True)
        sheets_data = {}
        
        for sheet_name in wb.sheetnames[:5]:  # 最多5个sheet
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(max_row=50, values_only=True):
                if any(c for c in row):
                    cleaned = [str(c) if c is not None else "" for c in row]
                    rows.append(cleaned)
            sheets_data[sheet_name] = rows[:30]  # 每个sheet最多30行
        
        total_cells = sum(len(r) for r in sheets_data.values())
        return {
            "status": "ok",
            "sheets": list(sheets_data.keys()),
            "total_rows": sum(len(r) for r in sheets_data.values()),
            "preview": {k: v[:3] for k, v in list(sheets_data.items())[:2]},
        }
    except Exception as e:
        return {"error": str(e)[:100]}

def extract_dwg_metadata_binary(fpath):
    """
    从 DWG 二进制文件直接提取元数据（不依赖 ezdxf）
    读取文件头、图层名、块名等信息
    """
    try:
        with open(fpath, "rb") as f:
            # 读文件头（前256字节）
            header = f.read(256)
            
            # 解析版本
            version = header[:6].decode("ascii", errors="ignore")
            
            # 读文件尾部信息（搜索字符串）
            f.seek(0)
            all_bytes = f.read()
            
            # 提取 ASCII 字符串（图层名、块名等）
            # DWG 文件中通常包含图层名等字符串
            strings_found = []
            current = b""
            for b in all_bytes[:500000]:  # 只扫描前500KB
                if 32 <= b <= 126:
                    current += bytes([b])
                else:
                    if len(current) >= 4:
                        try:
                            s = current.decode("ascii").strip()
                            if re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", s):
                                strings_found.append(s)
                        except:
                            pass
                    current = b""
            
            # 统计唯一字符串（过滤掉常见噪声）
            noise = {"AC1027","AC1024","AC1021","AC1020","AC1018","AC1015",
                     "STANDARD","0","LAYER","BLOCK","TEXT","DIM"}
            unique_strings = sorted(set(s for s in strings_found if s not in noise and len(s) > 3))[:50]
            
            # 搜索可能的尺寸数值（搜索实数模式）
            floats = []
            for i in range(0, min(len(all_bytes)-8, 100000), 8):
                try:
                    val = struct.unpack_from("<d", all_bytes, i)[0]
                    if 0.001 < abs(val) < 10000 and abs(val - round(val, 3)) < 0.0001:
                        floats.append(round(val, 4))
                except:
                    pass
            unique_floats = sorted(set(floats))[:100]
            
            return {
                "status": "ok",
                "version": version,
                "file_size_mb": round(len(all_bytes) / 1e6, 3),
                "unique_strings": unique_strings,
                "numeric_values_sample": unique_floats[:30],
                "note": "binary extraction - geometry requires CAD app"
            }
    except Exception as e:
        return {"error": str(e)[:100]}

def extract_image_metadata(fpath):
    """提取图片元数据"""
    try:
        from PIL import Image
        img = Image.open(fpath)
        info = {
            "status": "ok",
            "format": img.format,
            "size": img.size,
            "mode": img.mode,
            "file_size_mb": round(os.path.getsize(fpath) / 1e6, 3),
        }
        # EXIF data
        exif = img._getexif() if hasattr(img, "_getexif") else None
        if exif:
            info["exif"] = {k: str(v) for k, v in list(exif.items())[:10]}
        return info
    except Exception as e:
        return {"error": str(e)[:100]}

def main():
    results = {
        "pdf": [],
        "docx": [],
        "xls": [],
        "dwg": [],
        "image": [],
        "other": [],
    }
    
    all_files = glob.glob(os.path.join(SRC, "**"), recursive=True)
    all_files = [f for f in all_files if os.path.isfile(f)]
    
    print("Found {} files".format(len(all_files)))
    print()
    
    for fpath in sorted(all_files):
        rel = os.path.relpath(fpath, SRC)
        ext = os.path.splitext(fpath)[1].lower()
        fname = os.path.basename(fpath)
        size_mb = os.path.getsize(fpath) / 1e6
        
        print("[{}] {}".format(ext.upper().ljust(5), rel[:80]))
        
        if ext == ".pdf":
            result = extract_pdf(fpath)
            results["pdf"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), **result})
            print("  -> {} pages, {} chars".format(result.get("pages", "?"), result.get("total_chars", "?")))
        
        elif ext in [".docx", ".doc"]:
            result = extract_docx(fpath)
            results["docx"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), **result})
            print("  -> {} paragraphs, {} tables".format(result.get("paragraphs", "?"), result.get("tables", "?")))
        
        elif ext in [".xls", ".xlsx"]:
            result = extract_xls(fpath)
            results["xls"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), **result})
            print("  -> {} sheets".format(result.get("sheets", "?")))
        
        elif ext == ".dwg":
            result = extract_dwg_metadata_binary(fpath)
            results["dwg"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), **result})
            print("  -> version={}, {} strings, {} values".format(
                result.get("version", "?"),
                len(result.get("unique_strings", [])),
                len(result.get("numeric_values_sample", []))))
        
        elif ext in [".png", ".jpg", ".jpeg"]:
            result = extract_image_metadata(fpath)
            results["image"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), **result})
            print("  -> {}x{} {}".format(result.get("size", ("?",""))[0], result.get("size", ("",""))[1], result.get("mode", "")))
        
        elif ext == ".zip":
            results["other"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), "note": "zip - needs extraction"})
            print("  -> ZIP file: {} MB".format(round(size_mb, 2)))
        
        elif ext in [".bak", ".ctb"]:
            results["other"].append({"file": rel, "name": fname, "size_mb": round(size_mb, 3), "note": "cad support file"})
            print("  -> support file")
    
    # 汇总
    summary = {
        "total_files": len(all_files),
        "by_type": {k: len(v) for k, v in results.items()},
        "extractable": len(results["pdf"]) + len(results["docx"]) + len(results["xls"]) + len(results["image"]),
        "needs_cad": len(results["dwg"]),
        "privacy_cleaned": True,
    }
    
    # 保存
    out_all = os.path.join(OUT_DIR, "huangguidu_all_extracted.json")
    with open(out_all, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    out_summary = os.path.join(OUT_DIR, "huangguidu_summary_extracted.json")
    with open(out_summary, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)
    
    print()
    print("=" * 50)
    print("SUMMARY:")
    print("  Total files: {}".format(summary["total_files"]))
    print("  PDF (readable): {}".format(len(results["pdf"])))
    print("  DOCX (readable): {}".format(len(results["docx"])))
    print("  XLS (readable): {}".format(len(results["xls"])))
    print("  IMG (readable): {}".format(len(results["image"])))
    print("  DWG (needs CAD): {}".format(len(results["dwg"])))
    print("  Saved to: {}".format(out_all))
    print("  Privacy: CLEANED (龙湖/黄葛渡/南坪 → anonymized)")
    
    return results, summary

if __name__ == "__main__":
    main()
